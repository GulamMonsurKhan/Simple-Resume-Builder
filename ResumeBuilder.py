from docx import Document
from docx.shared import Inches

document = Document()

# Profile picture
document.add_picture('Shrek.png', width = Inches(2.0))

#Name phone number and email details
name = input('What is your name? ')
phoneNumber = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name + ' | ' + phoneNumber + ' | ' + email)

# About me
document.add_heading('About me')
document.add_paragraph(input('Tell me about yourself? '))

# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No? ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic

        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')

while True:
    has_skills = input('Do you have any skills? Yes or No? ')
    if has_skills.lower() == 'yes':
        p = document.add_paragraph()
        
        skills = input('Enter your skills ')

        p.add_run(skills + ' ')
        p.style = 'List Bullet'
    else:
        break


document.save(f"{name}_Resume.docx")